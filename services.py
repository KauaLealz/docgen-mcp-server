import os
import sys
import tempfile
import uuid
import zipfile
import csv
import re
import io
import difflib
import json
import xml.etree.ElementTree as ET
from typing import Optional, List, Dict, Any

import docx
import pypdf
import openpyxl
import mammoth
import markdownify
from bs4 import BeautifulSoup
from playwright.sync_api import sync_playwright
from reportlab.pdfgen import canvas
from reportlab.lib.colors import Color
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

from security import validate_path, validate_write_path
from config import get_read_sheet_max_rows, get_scan_max_matches


def odt_xml_to_plain(xml_bytes: bytes) -> str:
    root = ET.fromstring(xml_bytes)
    texts = []
    for elem in root.iter():
        if elem.text and elem.text.strip():
            texts.append(elem.text.strip())
    return "\n\n".join(texts)

def read_odt(path_resolved: str) -> str:
    with zipfile.ZipFile(path_resolved, "r") as z:
        try:
            content = z.read("content.xml")
        except KeyError:
            raise ValueError("content.xml nao encontrado no ODT.")
    return odt_xml_to_plain(content)

def read_ods_to_rows(ods_path: str, sheet_name: Optional[str] = None) -> List[List[Any]]:
    ns = {
        "table": "urn:oasis:names:tc:opendocument:xmlns:table:1.0",
        "text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0"
    }
    with zipfile.ZipFile(ods_path, "r") as z:
        content = z.read("content.xml")
    root = ET.fromstring(content)
    tables = root.findall(".//table:table", ns)
    if not tables:
        return []
    target_table = tables[0]
    if sheet_name:
        for t in tables:
            name_attr = f"{{{ns['table']}}}name"
            if t.attrib.get(name_attr) == sheet_name:
                target_table = t
                break
    rows = []
    for row_elem in target_table.findall("table:table-row", ns):
        row_repeat_attr = f"{{{ns['table']}}}number-rows-repeated"
        row_repeat = row_elem.attrib.get(row_repeat_attr)
        r_rep = int(row_repeat) if row_repeat and row_repeat.isdigit() else 1
        if r_rep > 1000:
            r_rep = 1000
        row = []
        cells = row_elem.findall("table:table-cell", ns)
        for cell in cells:
            repeat_attr = f"{{{ns['table']}}}number-columns-repeated"
            col_repeat = cell.attrib.get(repeat_attr)
            c_rep = int(col_repeat) if col_repeat and col_repeat.isdigit() else 1
            if c_rep > 1000:
                c_rep = 1000
            val_p = cell.find("text:p", ns)
            val = val_p.text if val_p is not None else ""
            for _ in range(c_rep):
                row.append(val)
        while row and row[-1] == "":
            row.pop()
        for _ in range(r_rep):
            rows.append(list(row))
    while rows and not any(cell != "" for cell in rows[-1]):
        rows.pop()
    return rows

def truncate_markdown(text: str, max_chars: Optional[int], preview_only: Optional[bool]) -> tuple:
    preview_cap = 12000
    cap = max_chars if max_chars is not None else (preview_cap if preview_only else None)
    if cap is None or len(text) <= cap:
        return text, False
    return f"{text[:cap]}\n\n_[… texto truncado; use maxChars ou desative previewOnly …]_", True

def read_doc_service(file_path: str, options: dict) -> dict:
    resolved = validate_path(file_path, True)
    ext = os.path.splitext(resolved)[1].lower()
    
    note = None
    if options.get("includeComments"):
        note = "includeComments: comentarios em Word/PDF podem nao ser extraidos integralmente nesta versao."
        
    if ext == ".pdf":
        reader = pypdf.PdfReader(resolved)
        pages_text = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                pages_text.append(t)
        md = "\n".join(pages_text).strip() if pages_text else "(PDF sem texto extraivel.)"
        md, truncated = truncate_markdown(md, options.get("maxChars"), options.get("previewOnly"))
        return {"markdown": md, "imagePaths": [], "note": note, "truncated": truncated}
        
    if ext == ".odt":
        text = read_odt(resolved)
        md = text if text else "(ODT vazio.)"
        md, truncated = truncate_markdown(md, options.get("maxChars"), options.get("previewOnly"))
        return {"markdown": md, "imagePaths": [], "note": note, "truncated": truncated}
        
    if ext == ".docx":
        image_paths = []
        if not options.get("extractImages"):
            doc = docx.Document(resolved)
            md = "\n".join(p.text for p in doc.paragraphs).strip()
            md = md if md else "(DOCX sem texto.)"
            md, truncated = truncate_markdown(md, options.get("maxChars"), options.get("previewOnly"))
            return {"markdown": md, "imagePaths": [], "note": note, "truncated": truncated}
            
        tmp_dir = os.path.join(tempfile.gettempdir(), f"docgen-img-{uuid.uuid4().hex}")
        os.makedirs(tmp_dir, exist_ok=True)
        
        def convert_image(image):
            with image.open() as image_bytes:
                ext_img = ".png" if "png" in image.content_type else ".jpg" if "jpeg" in image.content_type else ".bin"
                name = f"image-{len(image_paths)+1}{ext_img}"
                out = os.path.join(tmp_dir, name)
                with open(out, "wb") as f:
                    f.write(image_bytes.read())
                image_paths.append(out)
                return {"src": out}
                
        with open(resolved, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file, convert_image=mammoth.images.inline(convert_image))
            html = result.value
            
        md = markdownify.markdownify(html, heading_style="ATX").strip()
        md = md if md else "(DOCX sem texto.)"
        md, truncated = truncate_markdown(md, options.get("maxChars"), options.get("previewOnly"))
        return {"markdown": md, "imagePaths": image_paths, "note": note, "truncated": truncated}
        
def parse_range(range_str: str, max_row: int, max_col: int):
    if not range_str or not range_str.strip():
        return None
    m = re.match(r"^([A-Z]+)(\d+):([A-Z]+)(\d+)$", range_str.strip(), re.IGNORECASE)
    if not m:
        return None
    def col_to_idx(letters):
        n = 0
        for ch in letters.upper():
            n = n * 26 + (ord(ch) - 64)
        return n - 1
    c0 = col_to_idx(m.group(1))
    r0 = int(m.group(2)) - 1
    c1 = col_to_idx(m.group(3))
    r1 = int(m.group(4)) - 1
    return {
        "r0": max(0, r0),
        "r1": min(max_row - 1, r1),
        "c0": max(0, c0),
        "c1": min(max_col - 1, c1)
    }

def read_sheet_service(file_path: str, options: dict) -> dict:
    resolved = validate_path(file_path, True)
    ext = os.path.splitext(resolved)[1].lower()
    
    if ext == ".csv":
        with open(resolved, mode="r", encoding="utf-8-sig") as f:
            reader = csv.reader(f)
            rows = list(reader)
    elif ext == ".ods":
        rows = read_ods_to_rows(resolved, options.get("sheetName"))
    elif ext in {".xlsx", ".xlsm", ".xls"}:
        wb = openpyxl.load_workbook(resolved, read_only=True, data_only=True)
        name = options.get("sheetName")
        if not name or name not in wb.sheetnames:
            name = wb.sheetnames[0]
        ws = wb[name]
        rows = [[cell for cell in row] for row in ws.iter_rows(values_only=True)]
    else:
        raise ValueError(f"Extensao nao suportada para read_sheet: {ext}.")
        
    if not rows:
        return {"asJson": options.get("asJson"), "markdown": "", "jsonRows": [], "truncated": False, "rowLimit": 0}
        
    max_row = len(rows)
    max_col = max(len(r) for r in rows) if rows else 0
    for r in rows:
        if len(r) < max_col:
            r.extend([""] * (max_col - len(r)))
            
    range_str = options.get("range")
    if range_str and range_str.strip():
        bounds = parse_range(range_str, max_row, max_col)
        if not bounds:
            raise ValueError(f"range invalido: \"{range_str}\". Use o formato Excel com colunas e linhas, ex.: A1:D10 ou B2:B50.")
        sliced_rows = []
        for r_idx in range(bounds["r0"], bounds["r1"] + 1):
            row = rows[r_idx]
            sliced_rows.append(row[bounds["c0"]:bounds["c1"] + 1])
        rows = sliced_rows
        
    env_cap = get_read_sheet_max_rows()
    max_rows = options.get("maxRows")
    cap = min(max_rows, env_cap) if (max_rows is not None and max_rows > 0) else (min(500, env_cap) if options.get("previewOnly") else env_cap)
    
    max_rows_total = 1 + cap
    truncated = len(rows) > max_rows_total
    sliced = rows[:max_rows_total] if truncated else rows
    
    if options.get("asJson"):
        headers = [str(h).strip() or f"col_{i}" for i, h in enumerate(sliced[0])]
        out_rows = []
        for r in sliced[1:]:
            obj = {}
            for idx, h in enumerate(headers):
                val = r[idx] if idx < len(r) else ""
                obj[h] = val
            out_rows.append(obj)
        return {"asJson": True, "jsonRows": out_rows, "truncated": truncated, "rowLimit": cap}
        
    esc = lambda c: str(c if c is not None else "").replace("|", "\\|").replace("\n", " ")
    header_line = "| " + " | ".join(esc(h) for h in sliced[0]) + " |"
    sep_line = "| " + " | ".join("---" for _ in sliced[0]) + " |"
    body_lines = ["| " + " | ".join(esc(val) for val in r) + " |" for r in sliced[1:]]
    md = header_line + "\n" + sep_line + "\n" + "\n".join(body_lines)
    if truncated:
        md += f"\n\n_[Saida truncada: no maximo {cap} linhas de dados. Ajuste maxRows ou previewOnly._"
    return {"asJson": False, "markdown": md, "truncated": truncated, "rowLimit": cap}

def glob_to_regex(pattern: str) -> re.Pattern:
    p = pattern
    for ch in [".", "+", "^", "$", "{", "}", "(", ")", "|", "[", "]", "\\"]:
        p = p.replace(ch, "\\" + ch)
    p = p.replace("**", "{{GLOBSTAR}}")
    p = p.replace("*", "[^/]*")
    p = p.replace("?", ".")
    p = p.replace("{{GLOBSTAR}}", ".*")
    return re.compile(f"^{p}$", re.IGNORECASE)

def read_archive_service(file_path: str, pattern: Optional[str] = None) -> str:
    resolved = validate_path(file_path, True)
    ext = os.path.splitext(resolved)[1].lower()
    if ext != ".zip":
        raise ValueError(f"read_archive suporta apenas .zip. Recebido: {ext}")
        
    with zipfile.ZipFile(resolved, "r") as z:
        names = z.namelist()
        
    if pattern:
        pat = pattern.strip()
        rx = glob_to_regex(pat)
        names = [n for n in names if rx.match(n.replace("\\", "/"))]
        
    root = {}
    for name in sorted(names):
        parts = [p for p in name.replace("\\", "/").split("/") if p]
        cur = root
        for idx, p in enumerate(parts):
            if idx == len(parts) - 1 and not name.endswith("/"):
                cur[p] = "(file)"
            else:
                if p not in cur or not isinstance(cur[p], dict):
                    cur[p] = {}
                cur = cur[p]
                
    lines = []
    def walk(obj: dict, prefix=""):
        for k in sorted(obj.keys()):
            v = obj[k]
            is_file = v == "(file)"
            lines.append(f"{prefix}{k}" + ("" if is_file else "/"))
            if not is_file and isinstance(v, dict):
                walk(v, f"{prefix}{k}/")
                
    walk(root)
    tree = "\n".join(lines) if lines else "(vazio)"
    header = f"Arquivo: {resolved}\nEntradas{f' (filtro: {pattern})' if pattern else ''}: {len(names)}\n\n"
    return header + tree

def markdown_to_docx_paragraphs(doc, md_content: str):
    lines = md_content.replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(lines):
        raw = lines[i]
        t = raw.strip()
        if not t:
            i += 1
            continue
        if t.startswith("```"):
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            for line in code:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = "Consolas"
                p.paragraph_format.space_before = docx.shared.Pt(3)
                p.paragraph_format.space_after = docx.shared.Pt(3)
            continue
        hm = re.match(r"^(#{1,6})\s+(.*)$", t)
        if hm:
            level = len(hm.group(1))
            text = hm.group(2).strip()
            doc.add_heading(text, level=min(level, 3))
            i += 1
            continue
        if re.match(r"^[\-\*]\s+", t):
            items = []
            while i < len(lines):
                lt = lines[i].strip()
                im = re.match(r"^[\-\*]\s+(.*)$", lt)
                if not im:
                    break
                items.append(im.group(1).strip())
                i += 1
            for item in items:
                doc.add_paragraph(item, style="List Bullet")
            continue
        para_lines = [raw]
        i += 1
        while i < len(lines) and lines[i].strip():
            nt = lines[i].strip()
            if re.match(r"^(#{1,6}\s|```|[\-\*]\s)", nt):
                break
            para_lines.append(lines[i])
            i += 1
        body = "\n".join(para_lines).strip()
        if body:
            for pl in body.split("\n"):
                doc.add_paragraph(pl)

def markdown_to_pdf_plain_lines(md: str) -> list:
    lines = md.replace("\r\n", "\n").split("\n")
    out = []
    i = 0
    while i < len(lines):
        t = lines[i].strip()
        if not t:
            out.append("")
            i += 1
            continue
        if t.startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                out.append(f"  {lines[i]}")
                i += 1
            if i < len(lines):
                i += 1
            continue
        hm = re.match(r"^(#{1,6})\s+(.*)$", t)
        if hm:
            level = len(hm.group(1))
            text = hm.group(2).strip()
            out.append(f"{'#' * min(level, 3)} {text}")
            i += 1
            continue
        if re.match(r"^[\-\*]\s+", t):
            while i < len(lines):
                lt = lines[i].strip()
                im = re.match(r"^[\-\*]\s+(.*)$", lt)
                if not im:
                    break
                out.append(f"• {im.group(1)}")
                i += 1
            continue
        out.append(lines[i])
        i += 1
    return out

def write_docx_markdown(content: str, out_path: str):
    doc = docx.Document()
    markdown_to_docx_paragraphs(doc, content)
    doc.save(out_path)

def write_docx_plain(content: str, out_path: str):
    doc = docx.Document()
    blocks = re.split(r"\n\s*\n", content)
    for block in blocks:
        block_str = block.strip()
        if not block_str:
            continue
        lines = block_str.split("\n")
        head = lines[0].strip() if lines else ""
        rest = "\n".join(lines[1:]).strip() if len(lines) > 1 else ""
        if head.startswith("# "):
            doc.add_heading(head[2:], level=1)
            continue
        if head.startswith("## "):
            doc.add_heading(head[3:], level=2)
            continue
        text = f"{head}\n{rest}" if rest else head
        for line in text.split("\n"):
            doc.add_paragraph(line)
    doc.save(out_path)

def merge_docx_template(template_path: str, output_path: str, merge_fields: dict):
    with zipfile.ZipFile(template_path, "r") as jin:
        with zipfile.ZipFile(output_path, "w") as jout:
            for item in jin.infolist():
                data = jin.read(item.filename)
                if item.filename == "word/document.xml":
                    xml_content = data.decode("utf-8")
                    for k, v in merge_fields.items():
                        token = f"{{{{{k}}}}}"
                        xml_content = xml_content.replace(token, str(v if v is not None else ""))
                    data = xml_content.encode("utf-8")
                jout.writestr(item, data)

def write_pdf_service(content: str, out_path: str, format_markdown: bool = False):
    import html
    doc = SimpleDocTemplate(out_path, pagesize=letter)
    styles = getSampleStyleSheet()
    style = styles["Normal"]
    story = []
    lines = markdown_to_pdf_plain_lines(content) if format_markdown else content.split("\n")
    for line in lines:
        if line.strip():
            story.append(Paragraph(html.escape(line), style))
        else:
            story.append(Spacer(1, 10))
    doc.build(story)

def write_doc_service(params: dict) -> str:
    out_path = validate_write_path(params["path"])
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    fmt = params.get("contentFormat", "markdown")
    dtype = params["type"]
    
    if dtype == "docx":
        if params.get("templatePath"):
            tpl = validate_path(params["templatePath"], True)
            merge_fields = params.get("mergeFields", {})
            merge_docx_template(tpl, out_path, merge_fields)
            return f"DOCX gravado em {out_path} (template + mergeFields)." if merge_fields else f"DOCX gravado em {out_path} (copia do template)."
        if fmt == "markdown":
            write_docx_markdown(params["content"], out_path)
            return f"DOCX gravado em {out_path} (Markdown)."
        else:
            write_docx_plain(params["content"], out_path)
            return f"DOCX gravado em {out_path}."
            
    if dtype == "pdf":
        write_pdf_service(params["content"], out_path, format_markdown=(fmt == "markdown"))
        return f"PDF gravado em {out_path}."
        
    raise ValueError(f"Tipo nao suportado: {dtype}")

def write_sheet_service(params: dict) -> str:
    out_path = validate_write_path(params["path"])
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    ext = os.path.splitext(out_path)[1].lower()
    data = params["data"]
    columns = params.get("columns", {})
    
    if not isinstance(data, list):
        data = []
        
    if not data:
        if ext == ".csv":
            with open(out_path, "w", encoding="utf-8-sig") as f:
                f.write("")
            return f"CSV vazio gravado em {out_path}."
        elif ext == ".xlsx":
            wb = openpyxl.Workbook()
            wb.save(out_path)
            return f"Planilha vazia gravada em {out_path}."
        else:
            raise ValueError("write_sheet: use extensao .xlsx ou .csv.")
            
    first = data[0]
    if isinstance(first, dict):
        keys = list(first.keys())
    elif isinstance(first, list):
        keys = [f"col{i+1}" for i in range(len(first))]
    else:
        keys = ["value"]
        
    headers = [columns.get(k, k) for k in keys]
    
    row_values = []
    for r in data:
        if isinstance(r, dict):
            row_values.append([r.get(k) for k in keys])
        elif isinstance(r, list):
            row_values.append(r)
        else:
            row_values.append([r])
            
    if ext == ".csv":
        with open(out_path, mode="w", newline="", encoding="utf-8-sig") as f:
            writer = csv.writer(f)
            writer.writerow(headers)
            writer.writerows(row_values)
        return f"CSV gravado em {out_path} ({len(row_values)} linha(s))."
        
    if ext == ".xlsx":
        if params.get("append") and os.path.exists(out_path):
            wb = openpyxl.load_workbook(out_path)
            ws = wb.worksheets[0]
            for rv in row_values:
                ws.append(rv)
            wb.save(out_path)
            return f"Linhas anexadas em {out_path} ({len(row_values)} linha(s))."
            
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws.append(headers)
        for rv in row_values:
            ws.append(rv)
        if params.get("freezePanes"):
            ws.freeze_panes = "A2"
        wb.save(out_path)
        return f"Planilha gravada em {out_path} ({len(row_values)} linha(s) de dados + cabecalho)."
        
    raise ValueError("write_sheet: use extensao .xlsx ou .csv.")

def render_slide_service(params: dict) -> str:
    out_path = validate_write_path(params["outputPath"])
    aspect_ratio = params["aspectRatio"]
    w = 1920 if aspect_ratio == "16:9" else 1024
    h = 1080 if aspect_ratio == "16:9" else 768
    
    slide_css = f"""
    @page {{ size: {w}px {h}px; margin: 0; }}
    html, body {{ margin: 0; padding: 0; background: #fff; }}
    .slide {{
      width: {w}px;
      min-height: {h}px;
      box-sizing: border-box;
      page-break-after: always;
      page-break-inside: avoid;
    }}
    .slide:last-child {{ page-break-after: auto; }}
    {params['css']}
    """
    full_html = f"<!DOCTYPE html><html lang='pt-BR'><head><meta charset='utf-8'/><style>{slide_css}</style></head><body>{params['html']}</body></html>"
    
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        page.set_viewport_size({"width": w, "height": h})
        page.set_content(full_html)
        
        if params["format"] == "pdf":
            page.pdf(
                path=out_path,
                width=f"{w}px",
                height=f"{h}px",
                print_background=True,
                prefer_css_page_size=True
            )
            browser.close()
            return f"Slides exportados para PDF: {out_path}"
            
        slides = page.evaluate("() => Array.from(document.querySelectorAll('.slide')).map(el => el.outerHTML)")
        html_parts = slides if slides else [params["html"]]
        
        zip_path = out_path if out_path.lower().endswith(".zip") else f"{out_path}.zip"
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
            for i, chunk in enumerate(html_parts):
                doc = f"<!DOCTYPE html><html lang='pt-BR'><head><meta charset='utf-8'/><style>{slide_css}\n{params['css']}</style></head><body>{chunk}</body></html>"
                z.writestr(f"slide-{i+1:03d}.html", doc)
            z.writestr("styles-reference.css", slide_css + "\n" + params["css"])
            
        browser.close()
    return f"Slides exportados como ZIP: {zip_path}"

def render_page_service(params: dict) -> str:
    out_path = validate_write_path(params["outputPath"])
    toc_css = ""
    body_extra = ""
    if params.get("generateTOC"):
        toc_css = """
          #docgen-toc { page-break-after: always; font-family: system-ui, sans-serif; }
          #docgen-toc h2 { font-size: 14pt; }
          #docgen-toc nav { margin-top: 12px; }
          #docgen-toc a { color: inherit; text-decoration: none; }
        """
        body_extra = """
          <div id="docgen-toc"><h2>Indice</h2><nav id="toc-nav"></nav></div>
          <script>
            document.addEventListener('DOMContentLoaded', () => {
              const nav = document.getElementById('toc-nav');
              const hs = document.querySelectorAll('h1, h2, h3');
              hs.forEach((h, i) => {
                if (!h.id) h.id = 'heading-' + i;
                const a = document.createElement('a');
                a.href = '#' + h.id;
                a.textContent = h.textContent || '';
                const p = document.createElement('p');
                p.appendChild(a);
                nav.appendChild(p);
              });
            });
          </script>
        """
        
    margin = params.get("margins") or {}
    top = margin.get("top", "20mm")
    right = margin.get("right", "15mm")
    bottom = margin.get("bottom", "20mm")
    left = margin.get("left", "15mm")
    
    full_html = f"<!DOCTYPE html><html lang='pt-BR'><head><meta charset='utf-8'/><style>{toc_css}\n{params['css']}</style></head><body>{params['html'] + body_extra}</body></html>"
    
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        page.set_content(full_html)
        page.pdf(
            path=out_path,
            format="A4",
            print_background=True,
            margin={"top": top, "right": right, "bottom": bottom, "left": left},
            display_header_footer=False
        )
        browser.close()
    return f"Documento PDF gerado: {out_path}"

def patch_doc_service(params: dict) -> str:
    action = params["action"]
    
    if action == "merge":
        out_path = validate_write_path(params["path"])
        os.makedirs(os.path.dirname(out_path), exist_ok=True)
        sources = [validate_path(s, True) for s in params.get("payload", {}).get("sources", [])]
        if not sources:
            raise ValueError("merge: informe payload.sources (PDFs na ordem desejada).")
            
        merger = pypdf.PdfMerger()
        for src in sources:
            if os.path.splitext(src)[1].lower() != ".pdf":
                raise ValueError(f"merge: apenas PDF suportado. Arquivo invalido: {src}")
            merger.append(src)
        merger.write(out_path)
        merger.close()
        return f"PDF mesclado gravado em {out_path} ({len(sources)} arquivo(s))."
        
    resolved = validate_path(params["path"], True)
    out_path = validate_write_path(resolved)
    ext = os.path.splitext(resolved)[1].lower()
    
    if action == "split":
        if ext != ".pdf":
            raise ValueError("split: apenas arquivo .pdf na entrada.")
        p = params.get("payload") or {}
        out_dir = p.get("outputDir")
        if out_dir:
            out_dir = validate_write_path(out_dir)
        else:
            base = os.path.splitext(resolved)[0]
            out_dir = os.path.abspath(f"{base}_split")
        os.makedirs(out_dir, exist_ok=True)
        
        reader = pypdf.PdfReader(resolved)
        count = len(reader.pages)
        for i, page in enumerate(reader.pages):
            writer = pypdf.PdfWriter()
            writer.add_page(page)
            part_path = os.path.join(out_dir, f"page-{i+1:03d}.pdf")
            with open(part_path, "wb") as f:
                writer.write(f)
        return f"PDF dividido em {count} arquivo(s) em {out_dir}."
        
    if action == "watermark":
        if ext != ".pdf":
            raise ValueError("watermark: apenas .pdf suportado nesta versao.")
        p = params.get("payload") or {}
        text = p.get("text", "CONFIDENCIAL")
        opacity = p.get("opacity", 0.15)
        
        reader = pypdf.PdfReader(resolved)
        writer = pypdf.PdfWriter()
        
        for page in reader.pages:
            width = float(page.mediabox.width)
            height = float(page.mediabox.height)
            
            packet = io.BytesIO()
            can = canvas.Canvas(packet, pagesize=(width, height))
            can.setFillColor(Color(0.6, 0.6, 0.6, alpha=opacity))
            can.setFont("Helvetica-Bold", min(width, height) / 18)
            can.drawCentredString(width / 2, height / 2, text)
            can.save()
            
            packet.seek(0)
            watermark_reader = pypdf.PdfReader(packet)
            page.merge_page(watermark_reader.pages[0])
            writer.add_page(page)
            
        with open(out_path, "wb") as f:
            writer.write(f)
        return f"Marca d'agua aplicada em {out_path}."
        
    if action == "replace_text":
        p = params.get("payload") or {}
        reps = p.get("replacements", [])
        if not reps:
            raise ValueError("replace_text: informe payload.replacements.")
            
        if ext == ".docx":
            temp_path = out_path + ".tmp"
            with zipfile.ZipFile(resolved, "r") as jin:
                with zipfile.ZipFile(temp_path, "w") as jout:
                    for item in jin.infolist():
                        data = jin.read(item.filename)
                        if item.filename == "word/document.xml":
                            xml_content = data.decode("utf-8")
                            for r in reps:
                                xml_content = xml_content.replace(r["from"], r["to"])
                            data = xml_content.encode("utf-8")
                        jout.writestr(item, data)
            os.replace(temp_path, out_path)
            return f"Substituicoes aplicadas no DOCX em {out_path}."
            
        if ext == ".pdf":
            raise ValueError("replace_text em PDF nao e suportado. Exporte para DOCX.")
            
        raise ValueError(f"replace_text: extensao nao suportada: {ext}")
        
    raise ValueError(f"Acao desconhecida: {action}")

def patch_sheet_service(file_path: str, updates: list) -> str:
    resolved = validate_path(file_path, True)
    out_path = validate_write_path(resolved)
    ext = os.path.splitext(resolved)[1].lower()
    if ext != ".xlsx":
        raise ValueError("patch_sheet suporta apenas .xlsx nesta versao.")
    if not updates:
        raise ValueError("Informe ao menos uma atualizacao em updates.")
        
    wb = openpyxl.load_workbook(resolved)
    ws = wb.worksheets[0]
    
    for u in updates:
        cell_ref = u["cell"]
        cell = ws[cell_ref]
        cell.value = u["value"]
        style = u.get("style")
        if style and isinstance(style, dict):
            if "font" in style:
                font_opts = style["font"]
                cell.font = Font(**font_opts)
            if "fill" in style:
                fill_opts = style["fill"]
                cell.fill = PatternFill(**fill_opts)
            if "alignment" in style:
                align_opts = style["alignment"]
                cell.alignment = Alignment(**align_opts)
            if "border" in style:
                border_opts = style["border"]
                sides = {}
                for direction in ["left", "right", "top", "bottom"]:
                    if direction in border_opts:
                        d_opts = border_opts[direction]
                        b_style = d_opts.get("style")
                        color_opt = d_opts.get("color")
                        color_str = None
                        if isinstance(color_opt, dict):
                            color_str = color_opt.get("argb") or color_opt.get("theme")
                        elif isinstance(color_opt, str):
                            color_str = color_opt
                        sides[direction] = Side(border_style=b_style, color=color_str)
                cell.border = Border(**sides)
            if "numFmt" in style:
                cell.number_format = style["numFmt"]
                
    wb.save(out_path)
    return f"Planilha atualizada: {out_path} ({len(updates)} celula(s))."

def scan_zip(zip_path: str, query_re: re.Pattern) -> List[str]:
    hits = []
    with zipfile.ZipFile(zip_path, "r") as z:
        for name in z.namelist():
            clean_name = name.replace("\\", "/")
            if not clean_name.endswith("/") and query_re.search(clean_name):
                hits.append(f"{zip_path}::{clean_name}")
    return hits

def read_file_snippet(file_path: str, max_bytes: int = 8000) -> str:
    with open(file_path, "rb") as f:
        data = f.read(max_bytes)
        return data.decode("utf-8", errors="ignore")

def scan_dir_service(root_path: str, query: str, recursive: bool, max_matches: Optional[int] = None) -> dict:
    resolved = validate_path(root_path, True)
    try:
        re_flags = re.IGNORECASE | re.MULTILINE | re.DOTALL
        query_re = re.compile(query, re_flags)
    except Exception:
        raise ValueError("query nao e uma expressao regular valida.")
        
    st = os.stat(resolved)
    cap = max_matches if (max_matches is not None and max_matches > 0) else get_scan_max_matches()
    
    def trim(arr: list) -> tuple:
        if len(arr) <= cap:
            return arr, False
        return arr[:cap], True
        
    if os.path.isfile(resolved) and resolved.lower().endswith(".zip"):
        hits = scan_zip(resolved, query_re)
        trimmed, truncated = trim(hits)
        return {"matches": trimmed, "truncated": truncated}
        
    if os.path.isfile(resolved):
        text = read_file_snippet(resolved)
        lines = text.split("\n")
        hits = []
        for i, line in enumerate(lines):
            if query_re.search(line):
                hits.append(f"{resolved}:{i + 1}:{line[:200]}")
        trimmed, truncated = trim(hits)
        return {"matches": trimmed, "truncated": truncated}
        
    matches = []
    for root, dirs, files in os.walk(resolved):
        dirs[:] = [d for d in dirs if not d.startswith(".") and d != "__pycache__"]
        if not recursive and root != resolved:
            continue
        for file in files:
            if file.startswith("."):
                continue
            full_path = os.path.join(root, file)
            base = os.path.basename(full_path)
            if query_re.search(base):
                matches.append(full_path)
                continue
            if full_path.lower().endswith(".zip"):
                matches.extend(scan_zip(full_path, query_re))
                continue
            try:
                text = read_file_snippet(full_path, 256 * 1024)
                if query_re.search(text):
                    matches.append(full_path)
            except Exception:
                pass
        if not recursive:
            break
            
    trimmed, truncated = trim(matches)
    return {"matches": trimmed, "truncated": truncated}

def diff_file_service(path_a: str, path_b: str, mode: str) -> str:
    a = validate_path(path_a, True)
    b = validate_path(path_b, True)
    
    if mode == "data":
        ext_a = os.path.splitext(a)[1].lower()
        ext_b = os.path.splitext(b)[1].lower()
        if ext_a != ext_b:
            return "diff data: as extensoes devem coincidir para modo data."
            
        if ext_a in {".xlsx", ".xls", ".csv", ".ods"}:
            def read_data_to_str(p, ext):
                if ext == ".csv":
                    with open(p, mode="r", encoding="utf-8-sig") as f:
                        return json.dumps(list(csv.reader(f)), indent=2)
                elif ext == ".ods":
                    return json.dumps(read_ods_to_rows(p), indent=2)
                else:
                    wb = openpyxl.load_workbook(p, data_only=True)
                    res = {}
                    for name in wb.sheetnames:
                        ws = wb[name]
                        res[name] = [[c for c in row] for row in ws.iter_rows(values_only=True)]
                    return json.dumps(res, indent=2)
                    
            ja = read_data_to_str(a, ext_a)
            jb = read_data_to_str(b, ext_b)
            if ja == jb:
                return "Planilhas equivalentes (modo data)."
                
            diff = difflib.unified_diff(ja.splitlines(), jb.splitlines(), fromfile=a, tofile=b)
            out = list(diff)
            return "\n".join(out[:5000]) + ("\n... (truncado)" if len(out) > 5000 else "")
            
        return "Modo data: apenas planilhas tabulares (.xlsx, .csv, .ods) suportadas."
        
    try:
        with open(a, "r", encoding="utf-8") as fa:
            lines_a = fa.readlines()
        with open(b, "r", encoding="utf-8") as fb:
            lines_b = fb.readlines()
    except Exception:
        return "Modo text: um dos arquivos nao e texto UTF-8 legivel; use modo data ou converta antes."
        
    diff = difflib.unified_diff(lines_a, lines_b, fromfile=a, tofile=b)
    out = list(diff)
    return "".join(out[:8000]) + ("\n... (truncado)" if len(out) > 8000 else "")

def bundle_zip_service(files: List[str], output_name: str) -> str:
    if not files:
        raise ValueError("Informe ao menos um arquivo em files.")
    resolved = [validate_path(f, True) for f in files]
    out_path_raw = output_name if ("/" in output_name or "\\" in output_name) else os.path.abspath(output_name)
    out_path = validate_write_path(out_path_raw)
    
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as z:
        for f in resolved:
            z.write(f, os.path.basename(f))
            
    return f"Arquivo ZIP criado: {out_path} ({len(files)} arquivo(s))."
