from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path

from utils.security import validate_image_path, validate_path
from utils.file_utils import generate_output_path


def create_docx(title: str, sections: list[dict], output_path: str = "") -> str:
    """Generate a Word document from structured sections.

    Supported section types:
      - heading: {"type": "heading", "text": "...", "level": 1}
      - paragraph: {"type": "paragraph", "text": "...", "bold": false}
      - table: {"type": "table", "headers": [...], "rows": [[...]], "style": "Light Grid Accent 1"}
      - image: {"type": "image", "path": "...", "width_inches": 5.0, "caption": "..."}
      - code_block: {"type": "code_block", "code": "...", "language": "python"}
      - list: {"type": "list", "items": [...], "ordered": false}
      - page_break: {"type": "page_break"}

    Returns the absolute path of the generated file.
    """
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    if title:
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    for section in sections:
        _render_section(doc, section)

    dest = generate_output_path(output_path)
    doc.save(str(dest))
    return str(dest)


def read_docx(file_path: str) -> dict:
    """Extract text, tables, and image info from a Word document.

    Returns:
        {
            "text": str (full text),
            "paragraphs": [{"text": str, "style": str}],
            "tables": [{"headers": [...], "rows": [[...]]}],
            "images_count": int,
            "metadata": {"author": ..., "created": ..., ...}
        }
    """
    resolved = validate_path(file_path, must_exist=True)
    doc = Document(str(resolved))

    paragraphs = []
    full_text_parts = []
    for p in doc.paragraphs:
        paragraphs.append({"text": p.text, "style": p.style.name})
        if p.text.strip():
            full_text_parts.append(p.text)

    tables = []
    for table in doc.tables:
        rows_data = []
        headers = []
        for i, row in enumerate(table.rows):
            cells = [cell.text for cell in row.cells]
            if i == 0:
                headers = cells
            else:
                rows_data.append(cells)
        tables.append({"headers": headers, "rows": rows_data})

    images_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            images_count += 1

    core = doc.core_properties
    metadata = {
        "author": core.author or "",
        "created": core.created.isoformat() if core.created else "",
        "modified": core.modified.isoformat() if core.modified else "",
        "title": core.title or "",
        "subject": core.subject or "",
    }

    return {
        "text": "\n".join(full_text_parts),
        "paragraphs": paragraphs,
        "tables": tables,
        "images_count": images_count,
        "metadata": metadata,
    }


def _render_section(doc: Document, section: dict):
    section_type = section.get("type", "paragraph")

    if section_type == "heading":
        level = section.get("level", 1)
        doc.add_heading(section.get("text", ""), level=min(level, 4))

    elif section_type == "paragraph":
        p = doc.add_paragraph()
        text = section.get("text", "")
        if section.get("bold"):
            run = p.add_run(text)
            run.bold = True
        elif section.get("italic"):
            run = p.add_run(text)
            run.italic = True
        else:
            p.add_run(text)
        if section.get("alignment") == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    elif section_type == "table":
        headers = section.get("headers", [])
        rows = section.get("rows", [])
        style = section.get("style", "Light Grid Accent 1")
        if not headers:
            return
        num_cols = len(headers)
        table = doc.add_table(rows=1 + len(rows), cols=num_cols, style=style)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = str(h)
            cell.paragraphs[0].runs[0].bold = True if cell.paragraphs[0].runs else False

        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                if j < num_cols:
                    table.rows[i + 1].cells[j].text = str(val)

        doc.add_paragraph()

    elif section_type == "image":
        image_path = validate_image_path(section["path"])
        width = Inches(section.get("width_inches", 5.0))
        doc.add_picture(str(image_path), width=width)
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        caption = section.get("caption")
        if caption:
            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap_p.add_run(caption)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)

    elif section_type == "code_block":
        code = section.get("code", "")
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(code)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 30, 30)
        doc.add_paragraph()

    elif section_type == "list":
        items = section.get("items", [])
        ordered = section.get("ordered", False)
        style_name = 'List Number' if ordered else 'List Bullet'
        for item in items:
            doc.add_paragraph(str(item), style=style_name)

    elif section_type == "page_break":
        doc.add_page_break()
